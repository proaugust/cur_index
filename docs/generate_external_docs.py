# -*- coding: utf-8 -*-
"""生成 cur_index 对外介绍 Word（中文 / 日文）。运行：uv run --with python-docx python docs/generate_external_docs.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent


def _set_run_font(run, name_cn: str, size: int = 11, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name_cn
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name_cn)


def add_title(doc: Document, text: str, font: str):
    p = doc.add_heading(text, level=0)
    for run in p.runs:
        _set_run_font(run, font, 22, True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, text: str, level: int, font: str):
    p = doc.add_heading(text, level=level)
    size = {1: 16, 2: 14, 3: 12}.get(level, 12)
    for run in p.runs:
        _set_run_font(run, font, size, True)


def add_p(doc: Document, text: str, font: str, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, font, size)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullets(doc: Document, items: list[str], font: str):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, font, 11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font: str):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, font, 10, True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            _set_run_font(run, font, 10)
    doc.add_paragraph()


def build_zh() -> Document:
    font = "微软雅黑"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    add_title(doc, "cur_index AI 能力演示平台", font)
    add_p(doc, "对外技术介绍文档", font, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本说明：基于当前仓库实现整理，面向产品介绍与技术交流。")
    _set_run_font(run, font, 10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_h(doc, "1. 项目概述", 1, font)
    add_p(
        doc,
        "cur_index 是一套可交互的 AI 业务能力演示与后台管理一体化平台。"
        "后端基于 FastAPI，前端基于 Vue3 + Element Plus；数据层采用 PostgreSQL（含 pgvector），"
        "可选 Redis 缓存，大模型通过 OpenAI 兼容接口（默认 DeepSeek）调用。"
        "用户登录后可按权限体验文档检索增强生成（RAG）、投诉语义归类、客户洞察、"
        "多智能体博弈、智能路由、人脸考勤等模块，并配套 RBAC 权限与运维监控。",
        font,
    )
    add_p(doc, "平台技术主线可概括为两类能力：", font)
    add_bullets(
        doc,
        [
            "生成式 AI（LLM）：问答综合、会议纪要、意图路由、Agent 决策与解说等；",
            "表示学习与表格机器学习：中文 Embedding 检索/归类，LightGBM + SHAP + K-Means 风险洞察。",
        ],
        font,
    )

    add_h(doc, "2. 总体技术架构", 1, font)
    add_p(
        doc,
        "表现层（Vue3 管理台）通过 Axios 调用 FastAPI；接口层按 system / modules / ops 分域；"
        "服务层承载 RAG、投诉、Insight、Agent、考勤等业务逻辑；基础设施包括 PostgreSQL + pgvector、"
        "Redis（失败可降级内存缓存）、本地 Embedding 模型（BAAI/bge-small-zh-v1.5）以及外部 LLM API。"
        "生产环境可采用 Docker 多阶段构建，将前端静态资源与后端一并部署（如 Hugging Face Space）。",
        font,
    )
    add_table(
        doc,
        ["层级", "技术", "职责"],
        [
            ["前端", "Vue3 / TypeScript / Element Plus / Pinia / ECharts", "管理台交互与可视化"],
            ["API", "FastAPI / JWT / RBAC", "鉴权与业务接口"],
            ["AI 检索", "sentence-transformers / pgvector / 全文检索", "语义与混合检索"],
            ["LLM", "OpenAI 兼容 API / LangChain（Agent 演示）", "生成、路由、多步编排"],
            ["表格 ML", "LightGBM / SHAP / scikit-learn K-Means", "风险评分与可解释分析"],
            ["数据与缓存", "PostgreSQL / Redis", "持久化与统计缓存"],
        ],
        font,
    )

    add_h(doc, "3. 核心业务模块介绍", 1, font)

    add_h(doc, "3.1 RAG 检索：混合检索与融合排序", 2, font)
    add_p(
        doc,
        "业务知识库将文档切块后写入向量库，支持多种检索模式。"
        "纯向量模式基于 pgvector 余弦相似度召回；混合模式（hybrid）在向量召回之外，"
        "并行使用 PostgreSQL 全文检索，再对向量分、全文分以及标题/章节路径命中进行加权融合排序。"
        "需要完整回答时，由大模型综合多条原文片段生成可读说明，并保留来源切块，实现检索增强生成（RAG）。",
        font,
    )
    add_p(doc, "技术要点：", font)
    add_bullets(
        doc,
        [
            "Embedding：BAAI/bge-small-zh-v1.5（中文语义向量）；",
            "混合召回：语义相似 + 关键词全文互补；",
            "融合排序：规则加权（向量、全文、标题/路径结构特征）；可对同章节切块做 Parent 扩节以补全上下文；",
            "RAG 回答：LLM 综合 + 原文溯源。",
        ],
        font,
    )
    add_p(
        doc,
        "说明：界面中的 hybrid_rerank 与 hybrid 共用同一套融合逻辑，属于基于分数与结构特征的规则重排；"
        "若需神经重排（如 Cross-Encoder），可作为后续增强方向。",
        font,
    )

    add_h(doc, "3.2 投诉归类：语义向量自动归类", 2, font)
    add_p(
        doc,
        "投诉模块在统一的语义向量空间中完成自动归类：投诉类型由多种子句嵌入的均值向量表示类中心，"
        "投诉文本嵌入后与各类中心计算余弦相似度，达到阈值即自动归类。"
        "当相似度不足时，大模型仅负责建议新类型名称，并经名称向量去重后落库，避免类目膨胀。"
        "此外支持相似投诉语义检索，以及将自然语言统计问题解析为结构化过滤条件后走数据库聚合，"
        "实现「归类—检索—问数」一体，而不是用大模型直接贴标签或编造统计数字。",
        font,
    )
    add_bullets(
        doc,
        [
            "核心算法：中文 Embedding + 类中心近邻（阈值可配置）；",
            "LLM 角色：新类命名、自然语言问数解析；",
            "与 RAG 同源技术栈：同一套向量能力，侧重归类与运营统计。",
        ],
        font,
    )

    add_h(doc, "3.3 Customer Insight AI：可解释客户风险洞察", 2, font)
    add_p(
        doc,
        "Customer Insight 面向客户运营与流失风险管理，将行为与画像特征转化为可落地的风险清单与行动建议。"
        "夜间批处理完成特征工程后，使用 LightGBM 输出流失/风险概率并划分高中低档；"
        "对高风险客户用 SHAP（TreeExplainer）解释关键驱动因素（如投诉、满意度缺口等），"
        "驱动挽留动作建议；K-Means 用于人群分群，将样本不足的沉默客户映射到相似群体以补充标签与归因，"
        "且不篡改 LightGBM 原始分数。决策中心支持 WHAT-IF 仿真：调整关键特征后再次推理，观察风险变化。",
        font,
    )
    add_table(
        doc,
        ["技术", "对业务的产出"],
        [
            ["LightGBM", "风险分数、高/中/低档、可排序高风险清单"],
            ["SHAP", "中文特征贡献解释，支撑可解释挽留建议"],
            ["K-Means", "客户分群；沉默客户外推标签与簇级归因"],
            ["决策仿真", "干预假设下的风险变化（WHAT-IF）"],
        ],
        font,
    )
    add_p(
        doc,
        "该模块体现「可解释机器学习」在运营场景的闭环：打分 → 解释 → 建议 → 仿真，"
        "与纯对话式大模型能力形成互补。模型不可用时可回退规则打分，保证演示链路完整。",
        font,
    )

    add_h(doc, "3.4 多智能体博弈：炸金花（Agent 模拟对局）", 2, font)
    add_p(
        doc,
        "该模块将三名风格迥异的大模型玩家置于同一牌桌，演示「LLM 决策 + 硬规则可信执行」的多智能体架构。"
        "洗牌、牌型比较、注额计算、比牌与结算由确定性规则引擎完成；大模型只在人设与可见信息提示下，"
        "输出跟注、加注、比牌或弃牌等决策（含短思考）。系统会纠偏明显违规决策；模型不可用时按牌力规则代打。"
        "另有裁判模型根据公开行动生成直播风解说，禁止编造隐藏牌型。",
        font,
    )
    add_bullets(
        doc,
        [
            "人设示例：赌徒（激进诈唬）、老炮（稳健）、数学家（赔率思维）；",
            "架构原则：规则引擎算牌算账，LLM 只做人设决策与解说；",
            "工程保障：JSON 结构化决策、违规拦截、失败兜底，避免牌局卡死。",
        ],
        font,
    )

    add_h(doc, "3.5 智能路由：口语调用后台不同接口", 2, font)
    add_p(
        doc,
        "智能路由将用户一句话映射为业务意图（如天气、员工、邮件等），再分发到对应后端能力："
        "优先由大模型做意图分类，失败时回退关键词规则；命中天气则调用气象工具，"
        "命中员工则按工号查询考勤档案等接口。体现「自然语言作为统一入口、后端能力可插拔编排」的轻量 Agent 路由思路，"
        "适合讲解企业内多系统入口整合。",
        font,
    )

    add_h(doc, "3.6 人脸打卡：端侧识别与服务端比对", 2, font)
    add_p(
        doc,
        "考勤模块在浏览器用 face-api 完成人脸检测并提取 128 维描述子，上传后端后与人员库做欧氏距离 1:N 匹配；"
        "未命中可自动建档，命中则写入打卡记录。系统刻意限制特征随意融合，以降低误匹配污染。"
        "用于验证生物特征识别与业务打卡流程的端到端可行性（部署环境若剥离模型权重，则该能力可能不可用）。",
        font,
    )

    add_h(doc, "3.7 会议整理：大模型结构化改写", 2, font)
    add_p(
        doc,
        "用户粘贴口语化会议记录后，由大模型按「精简」或「正规」风格生成 Markdown 纪要。"
        "正规版强调总结、分议题要点、决议与待办，并约束基于原文、禁止编造。"
        "该模块技术路径简洁，重点展示 LLM 在办公场景的结构化信息抽取与改写能力。",
        font,
    )

    add_h(doc, "3.8 COBOL → Java：多 Agent 迁移方法论演示", 2, font)
    add_p(
        doc,
        "该模块以七步流水线（扫描、分类、知识入库、依赖成图、映射 Spring Boot、闭环校验、片段测试）"
        "展示遗留系统迁移的人机协同叙事与界面形态，用于验证方案与文档想法。"
        "当前实现以精心设计的示范数据驱动各步骤输出，不执行真实 COBOL 解析或代码生成，"
        "适合作为方法论与产品原型交流材料。",
        font,
    )

    add_h(doc, "4. AI 在各模块中的角色对照", 1, font)
    add_table(
        doc,
        ["模块", "AI / 算法角色", "关键技术关键词"],
        [
            ["RAG 知识库", "检索 + 可选生成综合回答", "Embedding / pgvector / Hybrid / RAG"],
            ["投诉归类", "向量归类；LLM 命名与问数解析", "语义相似度 / NL→SQL 参数"],
            ["Customer Insight", "风险打分、解释、分群、仿真", "LightGBM / SHAP / K-Means"],
            ["炸金花", "多智能体决策与解说", "Agent / Persona / 规则引擎"],
            ["智能路由", "意图识别 → 接口分发", "LLM Intent / Tool Routing"],
            ["人脸打卡", "人脸嵌入与匹配", "face-api / 欧氏距离"],
            ["会议整理", "单次结构化改写", "Prompt / Markdown 纪要"],
            ["COBOL 迁移", "多步骤演示叙事", "Multi-Agent 故事板（示范数据）"],
        ],
        font,
    )

    add_h(doc, "5. 总结", 1, font)
    add_p(
        doc,
        "cur_index 将检索增强生成、语义归类、可解释表格机器学习、多智能体博弈与自然语言接口路由等能力"
        "集中在同一管理台中演示，便于技术交流、方案验证与能力展示。"
        "对外介绍时建议突出三条主线：① Hybrid RAG 与可溯源问答；② Insight 的 LightGBM + SHAP 可解释风控；"
        "③ 炸金花所代表的「LLM 决策 + 硬规则可信执行」Agent 架构。",
        font,
    )

    return doc


def build_ja() -> Document:
    font = "游ゴシック"
    # Fallback widely available on Windows JP / CN
    font = "Yu Gothic"
    doc = Document()

    add_title(doc, "cur_index AI能力デモプラットフォーム", font)
    add_p(doc, "対外技術紹介ドキュメント", font, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版説明：現行リポジトリ実装に基づく製品・技術交流向け資料。")
    _set_run_font(run, font, 10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_h(doc, "1. プロジェクト概要", 1, font)
    add_p(
        doc,
        "cur_index は、対話可能な AI 業務能力デモとバックオフィス管理を一体化したプラットフォームです。"
        "バックエンドは FastAPI、フロントエンドは Vue3 + Element Plus を採用し、"
        "データ層は PostgreSQL（pgvector 含む）、任意で Redis キャッシュ、"
        "大規模言語モデルは OpenAI 互換 API（既定 DeepSeek）を利用します。"
        "ログイン後、権限に応じて RAG（検索拡張生成）、苦情の意味的分類、顧客インサイト、"
        "マルチエージェント対戦、インテントルーティング、顔認証打刻などを体験でき、"
        "RBAC と運用監視も備えています。",
        font,
    )
    add_p(doc, "技術の主軸は次の二系統に整理できます。", font)
    add_bullets(
        doc,
        [
            "生成 AI（LLM）：回答統合、議事録整理、意図ルーティング、Agent 意思決定と実況など；",
            "表現学習と表形式 ML：中国語 Embedding による検索・分類、LightGBM + SHAP + K-Means によるリスク洞察。",
        ],
        font,
    )

    add_h(doc, "2. 全体技術アーキテクチャ", 1, font)
    add_p(
        doc,
        "表現層（Vue3 管理画面）が Axios 経由で FastAPI を呼び出し、API 層は system / modules / ops に分割されます。"
        "サービス層が RAG・苦情・Insight・Agent・打刻などの業務を担い、基盤には PostgreSQL + pgvector、"
        "Redis（障害時はメモリキャッシュへフォールバック）、ローカル Embedding（BAAI/bge-small-zh-v1.5）、"
        "外部 LLM API があります。本番は Docker マルチステージ構築によりフロント静的資産とバックエンドを一体配備"
        "（例：Hugging Face Space）できます。",
        font,
    )
    add_table(
        doc,
        ["層", "技術", "役割"],
        [
            ["フロント", "Vue3 / TypeScript / Element Plus / Pinia / ECharts", "管理UIと可視化"],
            ["API", "FastAPI / JWT / RBAC", "認証認可と業務API"],
            ["AI検索", "sentence-transformers / pgvector / 全文検索", "意味・ハイブリッド検索"],
            ["LLM", "OpenAI互換API / LangChain（Agentデモ）", "生成・ルーティング・多段編成"],
            ["表形式ML", "LightGBM / SHAP / scikit-learn K-Means", "リスクスコアと説明可能分析"],
            ["データ", "PostgreSQL / Redis", "永続化と統計キャッシュ"],
        ],
        font,
    )

    add_h(doc, "3. 主要業務モジュール", 1, font)

    add_h(doc, "3.1 RAG検索：ハイブリッド検索と融合ランキング", 2, font)
    add_p(
        doc,
        "業務ナレッジベースは文書をチャンク化してベクトルストアへ格納し、複数の検索モードを提供します。"
        "ベクトルのみのモードは pgvector のコサイン類似度で召還し、ハイブリッド（hybrid）はベクトルに加え"
        "PostgreSQL 全文検索を並行実行し、ベクトル点・全文点・タイトル／章パス一致を重み付きで融合ソートします。"
        "完全な回答が必要な場合は、LLM が複数の原文断片を統合して可読な説明を生成し、出典チャンクを保持する"
        "検索拡張生成（RAG）を実現します。",
        font,
    )
    add_bullets(
        doc,
        [
            "Embedding：BAAI/bge-small-zh-v1.5（中国語意味ベクトル）；",
            "ハイブリッド召還：意味類似とキーワード全文の補完；",
            "融合ソート：規則的重み付け（ベクトル・全文・構造特徴）；同一セクションの Parent 拡張で文脈補完；",
            "RAG回答：LLM統合＋原文トレーサビリティ。",
        ],
        font,
    )
    add_p(
        doc,
        "注：画面上の hybrid_rerank は hybrid と同一の融合ロジックであり、スコアと構造特徴に基づく規則リランキングです。"
        "Cross-Encoder 等の神経リランキングは今後の強化候補です。",
        font,
    )

    add_h(doc, "3.2 苦情分類：意味ベクトルによる自動分類", 2, font)
    add_p(
        doc,
        "苦情モジュールは統一の意味ベクトル空間で自動分類を行います。"
        "カテゴリは複数シード文の埋め込み平均を中心ベクトルとし、苦情テキスト埋め込みとのコサイン類似度が"
        "閾値以上なら自動割当します。類似度不足時は LLM が新カテゴリ名の提案のみを担当し、"
        "名称ベクトルで重複排除して登録し、カテゴリ膨張を抑制します。"
        "類似苦情の意味検索や、自然言語の統計質問を構造化フィルタへ変換して DB 集計する機能もあり、"
        "「分類—検索—問数」を一体で提供します（LLM が直接ラベル付けや数値捏造を行う設計ではありません）。",
        font,
    )
    add_bullets(
        doc,
        [
            "中核：中国語 Embedding＋カテゴリ中心近傍（閾値設定可）；",
            "LLMの役割：新カテゴリ命名、自然言語問数の解析；",
            "RAGと同系スタック：同一ベクトル基盤を分類・運用統計に適用。",
        ],
        font,
    )

    add_h(doc, "3.3 Customer Insight AI：説明可能な顧客リスク洞察", 2, font)
    add_p(
        doc,
        "Customer Insight は顧客オペレーションと離反リスク管理向けに、行動・属性特徴を実務的なリスク一覧と"
        "アクション提案へ変換します。夜間バッチで特徴量を構築し、LightGBM で離反／リスク確率を算出し高中低に区分、"
        "高リスク顧客には SHAP（TreeExplainer）で主要要因（苦情、満足度ギャップ等）を説明し挽留提案を駆動します。"
        "K-Means は顧客クラスタリングに用い、サンプル不足のサイレント顧客を類似群へ写像してラベル・帰属を補完し、"
        "LightGBM の元スコアは改変しません。意思決定センターは WHAT-IF シミュレーション（特徴変更後の再推論）に対応します。",
        font,
    )
    add_table(
        doc,
        ["技術", "業務へのアウトプット"],
        [
            ["LightGBM", "リスクスコア、高／中／低、並べ替え可能な高リスク一覧"],
            ["SHAP", "中国語特徴寄与の説明、説明可能な挽留提案"],
            ["K-Means", "顧客クラスタ；サイレント顧客の外挿ラベルとクラスタ帰属"],
            ["意思決定シミュ", "介入仮説下のリスク変化（WHAT-IF）"],
        ],
        font,
    )
    add_p(
        doc,
        "本モジュールは「説明可能機械学習」のオペレーション閉ループ（スコア→説明→提案→シミュ）を示し、"
        "対話型 LLM 能力を補完します。モデル未準備時は規則スコアへフォールバックし、デモ経路を維持します。",
        font,
    )

    add_h(doc, "3.4 マルチエージェント対戦：炸金花（Agent による賭博シミュレーション）", 2, font)
    add_p(
        doc,
        "本モジュールは性格の異なる 3 名の LLM プレイヤーを同一卓に置き、"
        "「LLM 意思決定＋硬規則による信頼できる実行」のマルチエージェント構成を演示します。"
        "シャッフル、役比較、賭け金計算、勝負判定は決定論的ルールエンジンが担当し、"
        "LLM はペルソナと可視情報に基づきコール／レイズ／ショーダウン／フォールド等を JSON で出力します。"
        "明らかな违规はシステムが補正し、モデル障害時は手役強度ルールで代理プレイします。"
        "審判 LLM は公開アクションのみから実況風解説を生成し、隠れ役の捏造を禁止します。",
        font,
    )
    add_bullets(
        doc,
        [
            "ペルソナ例：ギャンブラー（攻撃的ブラフ）、ベテラン（堅実）、数学者（オッズ思考）；",
            "設計原則：ルールエンジンが札と会計を担当、LLM は意思決定と解説のみ；",
            "工学的担保：構造化出力、违规遮断、障害時フォールバックで卓の停止を防止。",
        ],
        font,
    )

    add_h(doc, "3.5 スマートルーティング：口語からバックエンド API 呼び出し", 2, font)
    add_p(
        doc,
        "スマートルーティングはユーザーの一文を業務インテント（天気・社員・メール等）へ写像し、"
        "対応するバックエンド能力へ振り分けます。優先的に LLM で意図分類し、失敗時はキーワード規則へフォールバック。"
        "天気なら気象ツール、社員なら社員番号で勤怠プロファイル API を呼びます。"
        "「自然言語を統一入口とし、バックエンド能力を差し替え可能な編成でつなぐ」軽量 Agent ルーティングの考え方を示します。",
        font,
    )

    add_h(doc, "3.6 顔認証打刻：端末側認識とサーバ側照合", 2, font)
    add_p(
        doc,
        "勤怠モジュールはブラウザの face-api で顔検出と 128 次元ディスクリプタ抽出を行い、"
        "サーバ側で人物庫とのユークリッド距離 1:N 照合を実施します。未一致時は自動登録、一致時は打刻記録を保存。"
        "特徴の安易な融合を制限し誤照合汚染を抑えます。"
        "生体認証と業務打刻の E2E 可行性検証向けです（配備環境でモデル重みを除外している場合は利用不可となり得ます）。",
        font,
    )

    add_h(doc, "3.7 会議整理：LLM による構造化リライト", 2, font)
    add_p(
        doc,
        "口語の会議メモを貼り付けると、LLM が「簡潔」または「正式」スタイルの Markdown 議事録を生成します。"
        "正式版は要約・議題別要点・決議・ToDo を重視し、原文に基づくこと・捏造禁止を制約します。"
        "技術経路は簡潔で、オフィスシーンにおける構造化情報抽出・リライト能力の演示に適します。",
        font,
    )

    add_h(doc, "3.8 COBOL → Java：マルチ Agent 移行方法論デモ", 2, font)
    add_p(
        doc,
        "本モジュールは 7 ステップ（スキャン、分類、知識登録、依存グラフ、Spring Boot 写像、閉ループ検証、断片テスト）"
        "でレガシー移行の人機協調ナラティブと UI を示し、方案・文書アイデアの検証に用います。"
        "現行実装は設計済みのデモデータで各ステップを駆動し、実 COBOL 解析やコード生成は行いません。"
        "方法論およびプロトタイプ交流向けの資料として位置づけます。",
        font,
    )

    add_h(doc, "4. 各モジュールにおける AI の役割対照", 1, font)
    add_table(
        doc,
        ["モジュール", "AI／アルゴリズムの役割", "キーワード"],
        [
            ["RAGナレッジ", "検索＋任意で生成統合回答", "Embedding / pgvector / Hybrid / RAG"],
            ["苦情分類", "ベクトル分類；LLM命名と問数解析", "意味類似 / NL→SQLパラメータ"],
            ["Customer Insight", "リスク採点・説明・クラスタ・シミュ", "LightGBM / SHAP / K-Means"],
            ["炸金花", "マルチエージェント意思決定と実況", "Agent / Persona / ルールエンジン"],
            ["スマートルート", "意図認識→API振分", "LLM Intent / Tool Routing"],
            ["顔認証打刻", "顔埋め込みと照合", "face-api / ユークリッド距離"],
            ["会議整理", "単発の構造化リライト", "Prompt / Markdown議事録"],
            ["COBOL移行", "多段デモナラティブ", "Multi-Agent ストーリーボード（デモデータ）"],
        ],
        font,
    )

    add_h(doc, "5. まとめ", 1, font)
    add_p(
        doc,
        "cur_index は、検索拡張生成、意味的分類、説明可能な表形式機械学習、マルチエージェント対戦、"
        "自然言語インタフェースルーティングなどを同一管理画面で演示し、技術交流・方案検証・能力展示を容易にします。"
        "対外紹介では次の三本柱を推奨します：① Hybrid RAG と出典付き回答；"
        "② Insight の LightGBM + SHAP による説明可能リスク管理；"
        "③ 炸金花に代表される「LLM 意思決定＋硬規則の信頼実行」Agent アーキテクチャ。",
        font,
    )

    return doc


def main():
    # ASCII filenames avoid Windows console/codepage garbling
    zh_path = OUT_DIR / "cur_index_external_intro_zh-CN.docx"
    ja_path = OUT_DIR / "cur_index_external_intro_ja-JP.docx"
    for old in OUT_DIR.glob("cur_index_*.docx"):
        try:
            old.unlink()
        except OSError:
            pass
    build_zh().save(zh_path)
    build_ja().save(ja_path)
    print(f"Wrote: {zh_path}")
    print(f"Wrote: {ja_path}")


if __name__ == "__main__":
    main()
